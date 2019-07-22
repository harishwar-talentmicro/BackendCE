try:
    import re
    from docx import Document
    from docx.shared import Inches
    import requests
    from io import BytesIO
    from io import StringIO
    import base64

    import argparse
    import datetime
    import pprint

    import os
    import sys
    from google.cloud import storage

    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.path.dirname(__file__)+"/cred.json"
    # os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "cred.json"

    # [START storage_upload_file]

    print (sys.argv)

    def docx_replace_regex(doc_obj, regex, replace):

        for p in doc_obj.paragraphs:
            if regex.search(p.text):
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if regex.search(inline[i].text):
                        text = regex.sub(replace, inline[i].text)
                        inline[i].text = text

        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_replace_regex(cell, regex, replace)

    # mobile_regex = re.compile(r"(\+)*[ 0-9\-]{8,14}")
    mobile_regex = re.compile(r"\(* *(\+)*[ \)0-9\-]{8,15}")
    # email_regex = re.compile(r"[a-zA-Z0-9\._]+@[a-zA-Z]+.[a-zA-Z]+")
    email_regex = re.compile(r"[a-zA-Z0-9._-]+@[a-zA-Z0-9._-]+\.[a-zA-Z0-9_-]+")
    replace_with = r""
    # "https://storage.googleapis.com/ezeone/6f8e5aaf-84dd-4ece-99c2-946cd469929f.docx"
    filename = sys.argv[1]+sys.argv[2]
    doc = Document(BytesIO(requests.get(filename).content))
    docx_replace_regex(doc, mobile_regex, replace_with)
    docx_replace_regex(doc, email_regex, replace_with)
    # doc.save('result2.docx')
    # with open("yourfile.ext", "rb") as image_file:
    #     encoded_string = base64.b64encode(image_file.read())
    section = doc.sections[0]
    header = section.header
    first_paragraph = header.paragraphs[0]
    image_paragraph = first_paragraph.insert_paragraph_before()
    header_image = image_paragraph.add_run()
    header_image.add_picture(
        BytesIO(requests.get(sys.argv[4]).content), width=Inches(1.5))
    text_paragraph = first_paragraph.insert_paragraph_before()
    header_text = text_paragraph.add_run()
    header_text.add_text(sys.argv[5])

   # paragraph.add_picture('download.png', width=Inches(1.25))
    try:
        storage_client = storage.Client()
        bucket = storage_client.get_bucket('ezeone')
        # '6f8e5aaf-84dd-4ece-99c2-946cd469929f12345.docx')
        blob = bucket.blob(sys.argv[3])
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        blob.upload_from_file(file_stream)
        print ('masked successfully')
    except Exception as e:
        print(e)
    # print("here")
    # with open(BytesIO(doc), "rb") as my_file:
    #     blob.upload_from_file(my_file)

       # print('File {} uploaded to {}.'.format(
       #     source_file_name,
       #     destination_blob_name))
    # [END storage_upload_file]
except Exception as e:
    print(e)
